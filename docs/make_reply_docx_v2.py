"""
Reply to Szafranski 2026-05-07 meeting summary.
Based on nife/dieckow_paper/dieckow_report_EN.tex and progress_report_hamilton_kegg.md.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# margins
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.8); s.right_margin = Cm(2.8)

NAVY = RGBColor(0x1F, 0x49, 0x7D)

def H(text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = NAVY
    return p

def P(text):
    return doc.add_paragraph(text)

def B(text):
    return doc.add_paragraph(text, style="List Bullet")

def bold_run(paragraph, bold_text, rest_text=""):
    p = doc.add_paragraph()
    r = p.add_run(bold_text); r.bold = True
    if rest_text:
        p.add_run(rest_text)
    return p

# ── header ──────────────────────────────────────────────
doc.add_heading("Reply: Progress Meeting 2026-05-07", 0)
meta = doc.add_paragraph()
for label, val in [("To: ", "Szafranski.Szymon@mh-hannover.de"),
                   ("From: ", "Keisuke Nishioka"),
                   ("Date: ", "2026-05-12"),
                   ("Subject: ", "Follow-up on guild dynamics modeling discussion")]:
    meta.add_run(label).bold = True
    meta.add_run(val + "\n")
doc.add_paragraph()

P("Dear Szymon,")
doc.add_paragraph()
P("Thank you for the detailed summary. Below I respond to each of your action items "
  "using the results from the current analysis.")
doc.add_paragraph()

# ── 1. What the model aims to achieve ───────────────────
H("1. What the model aims to achieve")
P("The model aims to test whether metabolic network constraints — derived from "
  "genome-scale metabolic models and curated experimental databases — can improve "
  "out-of-sample prediction of oral microbiome community dynamics.")
doc.add_paragraph()
P("Concretely: given week-1 community composition for a new patient, can we predict "
  "their week-2 and week-3 compositions using a shared ecological interaction matrix "
  "inferred from other patients?")
doc.add_paragraph()
bold_run(None, "Dataset: "); doc.paragraphs[-1].runs[0].bold = True  # fix below
p = doc.add_paragraph()
p.add_run("Dataset: ").bold = True
p.add_run("Dieckow et al. 2024 (npj Biofilms Microbiomes 10:85) — 10 patients (A–L), "
          "subgingival plaque on titanium implant abutments, 16S rRNA, weeks 1–3. "
          "ASVs aggregated to 10 class-level guilds following Szafrański et al. 2025 taxonomy.")
doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run("Dynamical model: ").bold = True
p2.add_run("Hamilton replicator ODE (0D limit of the Extended Hamilton Principle, "
           "Junker & Balzani 2021; Klempt et al. 2024). Governing equation:")
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.add_run("φ̇ᵢ = φᵢ ( bᵢ + Σⱼ Aᵢⱼ φⱼ − f̄(φ) )").italic = True
doc.add_paragraph()
P("where φᵢ is the relative abundance of guild i, A is the symmetric interaction matrix "
  "(A_ij = A_ji, from thermodynamic consistency of the Hamilton Principle), b_i is a "
  "patient-specific intrinsic growth rate, and f̄ = Σₖ φₖfₖ keeps Σφᵢ = 1 on the "
  "probability simplex. This is mathematically equivalent to the replicator equation "
  "(Taylor & Jonker 1978) with linear payoffs.")

# ── 2. What has been achieved ───────────────────────────
H("2. What has already been achieved")
B("Fitted the interaction matrix A on all 10 patients. Training RMSE = 0.0565, "
  "Pearson r = 0.951.")
B("Metabolic sign prior constructed from Szafrański 2025 Suppl. + eHOMD "
  "(experimentally curated PRODUCES/USES flows, L1 layer). "
  "An additional genome-scale metabolic modeling layer (L2) based on pFBA is being "
  "explored and is currently under validation.")
B("Leave-one-patient-out (LOO) cross-validation completed (10-fold). "
  "Best model (Hamilton + AGORA W=1.0):")

# results sub-table
tbl = doc.add_table(rows=6, cols=3)
tbl.style = "Table Grid"
for i, row in enumerate([
    ["Model", "LOO-RMSE", "LOO-BC (Bray-Curtis)"],
    ["gLV free (no prior)", "0.0588", "0.154"],
    ["Hamilton (no prior)", "0.0595", "—"],
    ["Hamilton + L1 (34 pairs)", "0.0516", "0.155"],
    ["Hamilton + L1+L2 / AGORA W=1.0", "0.0504", "0.147"],
    ["Persistence null", "—", "0.281"],
]):
    for j, v in enumerate(row):
        tbl.rows[i].cells[j].text = v
        if i == 0:
            tbl.rows[i].cells[j].paragraphs[0].runs[0].bold = True
doc.add_paragraph()

B("Sign agreement between the metabolic prior and the fitted A matrix: >90% of "
  "constrained pairs have the expected sign. The genome-scale modeling layer is "
  "being tested to assess whether it can further improve consistency.")
B("LOO A-matrix stability analysis: sign consistency of each guild-guild interaction "
  "assessed across all 10 LOO folds. Three regimes identified: interactions where "
  "data and prior agree strongly, interactions that are metabolically predicted but "
  "ecologically muted on the 3-week timescale, and data-driven interactions without "
  "metabolic annotation.")
B("External validation: the fitted A matrix (calibrated on Dieckow abutment data) "
  "was applied to 127 independent cross-sectional peri-implant samples from "
  "Joshi/Szafrański 2025. A Guild Dysbiosis Index (GDI, based on the ratio of "
  "dysbiotic vs. commensal guilds at equilibrium) correctly orders clinical severity "
  "(Kruskal-Wallis p < 0.0001, Spearman ρ = 0.37).")

# ── 3. Expected results ─────────────────────────────────
H("3. Which results were expected")
B("Metabolic sign priors were expected to regularise A and reduce overfitting → "
  "confirmed (LOO-RMSE −14% vs unconstrained gLV).")
B("Sign consistency between the experimental prior (eHOMD/Szafrański L1) and the "
  "fitted A was expected to be high, given well-documented cross-feeding relationships "
  "(e.g. Streptococcus → lactate → Veillonella) → confirmed: >90% of constrained "
  "pairs match the expected sign.")
B("Dominant commensal interactions — Actinobacteria ↔ Bacilli co-aggregation, "
  "Bacilli ↔ Betaproteobacteria syntrophy — were expected to be the strongest "
  "fitted interactions → confirmed (A = +1.61 and +1.83 respectively).")

# ── 4. Unexpected results ───────────────────────────────
H("4. Which results were unexpected")
B("The Bacilli ↔ Negativicutes (Streptococcus → lactate → Veillonella) interaction "
  "is one of the best-known oral cross-feeding syntrophies and has a strong "
  "experimental prior. However, the fitted A ≈ 0 — sign correct but ecologically "
  "muted on the 3-week abutment timescale. The most cited oral cross-feeding axis "
  "is apparently not the dominant ecological driver in this dataset.")
B("The fitted A matrix generalises across disease context: calibrated on caries "
  "recovery (abutment, 3 weeks), it correctly orders peri-implantitis dysbiosis "
  "in 127 independent samples. This suggests A captures fundamental ecological "
  "attractors rather than condition-specific dynamics.")
B("Mucositis ≈ Health at equilibrium (GDI −3.05 vs −2.90, p = 0.61). "
  "Clinically mucositis is an intermediate state, but the model places it in the "
  "commensal attractor basin — consistent with reversibility, but not with elevated "
  "inflammation markers.")
B("Patient A (CT2, high Actinobacteria, 3× mean) is a persistent outlier. "
  "This suggests a covariate-shift issue — Actinobacteria-dominant community types "
  "are underrepresented in the 10-patient training pool.")

# ── 5. Relation to data ─────────────────────────────────
H("5. How the model relates to available data")
p = doc.add_paragraph()
p.add_run("Calibration data: ").bold = True
p.add_run("Dieckow et al. 2024, 10 patients, weeks 1–3, 10 class-level guilds.")
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("External / clinical validation: ").bold = True
p.add_run("Joshi / Szafrański 2025 peri-implant cohort, N=127 cross-sectional samples "
          "(Health=56, Mucositis=39, PI=32). The same five genera (So, An, Vei, Fn, Pg) "
          "available in the Szafrański 5-genera subset map onto 5 of the 10 guilds, "
          "enabling direct comparison with the Hamilton attractor states.")
doc.add_paragraph()
P("The attractor principle used for validation: for the replicator equation, fixed points "
  "are determined by A alone (independent of b_i, confirmed empirically). "
  "Cross-sectional samples near disease equilibrium can therefore be projected onto "
  "the A-landscape inferred from longitudinal data.")

# ── 6. Assumptions ──────────────────────────────────────
H("6. Assumptions currently in the model")
B("Symmetric interaction matrix: A_ij = A_ji. "
  "Arises from the Hamilton Principle free energy (quadratic form). "
  "Reduces free parameters from 90 to 45 off-diagonal entries.")
B("Relative abundance dynamics on the probability simplex (Σφᵢ = 1). "
  "Appropriate for 16S amplicon data; avoids compositional bias of fitting "
  "absolute-abundance gLV to relative data.")
B("Patient-specific b_i captures inter-individual variation; "
  "shared A captures population-level guild ecology.")
B("Sign prior from eHOMD/Szafrański L1 (experimentally curated metabolite flows, "
  "34 constrained pairs). An additional metabolic modeling layer is being explored "
  "but is not yet finalized.")
B("10 guilds at SILVA class level (Szafrański taxonomy). "
  "Species-level interactions may cancel or reinforce at this aggregation level.")

# ── 7. Uncertain assumptions ────────────────────────────
H("7. Assumptions that remain uncertain")
B("Symmetric A excludes asymmetric interactions (commensalism, amensalism). "
  "Example: A secretes lactate → B benefits (+), B secretes H₂O₂ → A harmed (−) "
  "cannot be captured. N=10 is insufficient to identify asymmetric A (90 parameters).")
B("Condition-dependence of A: the current model uses a single A matrix. "
  "Your scenario (sucrose fermentation → lactate accumulation → low pH → "
  "inhibitory environment) requires A to depend on metabolite concentrations or pH. "
  "The data (3 timepoints) cannot identify a pH-dependent A(pH, t). "
  "An extension (A(t) = A_commensal + (A_dysbiotic − A_commensal) × sigmoid(pH − pH_crit)) "
  "would require explicit pH or metabolite measurements.")
B("Genome-scale metabolic modeling layer (L2): still under validation. "
  "The approach uses pFBA on guild-representative strains to predict cross-feeding signs, "
  "but the representative-strain choice, medium composition, and single-strain FBA "
  "approximation all introduce uncertainty. Detailed validation of this layer is ongoing.")
B("N=10 (marginal power): paired t-test AGORA vs L1-only gives p=0.08. "
  "Replication with N≥30 patients is needed for statistical significance.")
B("Stationarity assumption for external validation: peri-implant communities "
  "(cross-sectional) are assumed near ecological equilibrium. "
  "Unverified — chronic PI (months–years) supports this, but acute cases may not be.")

# ── 8. How DB interactions → parameters ─────────────────
H("8. How database-derived interactions become model parameters")
P("This is the step that needs clearest documentation. The workflow:")
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("Step 1 — Which pairs are constrained (topology)? ").bold = True
p.add_run("For each guild pair (i,j), the eHOMD/Szafrański supplementary provides "
          "explicit PRODUCES / USES / IS_INHIBITED_BY relationships at the metabolite level. "
          "If metabolite α is produced by guild j and consumed by guild i, this is a "
          "cross-feeding signal: net_flow[i,j] += w. Toxin secretion (H₂O₂, H₂S) "
          "gives net_flow[i,j] −= w. Pairs where net_flow = 0 are unconstrained.")
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("Step 2 — Sign: ").bold = True
p.add_run("sgn(F_ij) = sgn(net_flow[i,j] + net_flow[j,i]) / 2. "
          "A positive net flow means j benefits i more than it harms it → expected A_ij > 0. "
          "When evidence conflicts (e.g. a guild both produces a nutrient and a toxin affecting "
          "another guild), the net sum determines the expected sign. "
          "This is fully documented in the flow matrix F which can be inspected directly.")
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("Step 3 — Magnitude of A: ").bold = True
p.add_run("Not from databases. Magnitudes are inferred from fitting the ODE to "
          "the Dieckow time-series data via L-BFGS-B + JAX autodiff. "
          "The prior only constrains the sign (via a soft hinge penalty); "
          "the actual numerical A values are data-driven.")
doc.add_paragraph()

tbl2 = doc.add_table(rows=3, cols=3)
tbl2.style = "Table Grid"
for i, row in enumerate([
    ["Layer", "Source", "Status"],
    ["L1 (direct experimental)", "eHOMD + Szafrański 2025 Suppl. (PRODUCES/USES/IS_INHIBITED_BY)", "In use (34 pairs)"],
    ["L2 (genome-scale metabolic modeling)", "pFBA on guild-representative strains, oral-fluid medium", "Under validation"],
]):
    for j, v in enumerate(row):
        tbl2.rows[i].cells[j].text = v
        if i == 0:
            tbl2.rows[i].cells[j].paragraphs[0].runs[0].bold = True
doc.add_paragraph()

# ── 9. Lactate-sucrose ───────────────────────────────────
H("9. The sucrose–lactate–biofilm scenario")
P("The scenario you described maps onto an important limitation of the current model. "
  "The Hamilton ODE uses a time-invariant A matrix, so it cannot capture the regime "
  "shift where lactate accumulation and pH drop change the interaction from cross-feeding "
  "to inhibitory.")
doc.add_paragraph()
P("The relevant finding from the current analysis is that the Bacilli ↔ Negativicutes "
  "interaction (Streptococcus → lactate → Veillonella — the exact cross-feeding axis "
  "you described) is in the 'prior-constrained muted' regime: "
  "|F| = 5.0 (strongest AGORA signal among all pairs), but fitted A = +0.003 ≈ 0. "
  "Possible interpretations:")
B("The cross-feeding is real but operates on timescales longer than 3 weeks "
  "(muted in the abutment dataset).")
B("At the compositions in the Dieckow dataset (healthy recovery, not cariogenic), "
  "lactate concentrations may not accumulate enough to drive the feedback loop "
  "you described — exactly because the environment has not yet become low-pH.")
B("An A(t) extension with an explicit lactate/pH variable would be needed to "
  "distinguish these cases. This requires metabolite time-series data alongside 16S.")
doc.add_paragraph()
P("I will incorporate your lactate–sucrose sketches (once received via WhatsApp) into "
  "the design of the A(pH,t) extension.")

# ── 10. Compositional data / multivariate analysis ──────
H("10. Compositional data tables for multivariate analysis")
P("I will prepare the following tables:")
doc.add_paragraph()

tbl3 = doc.add_table(rows=4, cols=3)
tbl3.style = "Table Grid"
for i, row in enumerate([
    ["Dataset", "Contents", "Status"],
    ["Dieckow 2024 observed",
     "10 patients × 3 weeks × 10 guilds (relative abundances); "
     "metadata: patient ID, community type (CT1/CT2), week",
     "To prepare (from dieckow_otu/ data)"],
    ["Hamilton model predicted",
     "Same format as observed; LOO-predicted week 2/3 compositions "
     "for all 10 patients + full-training predictions",
     "To prepare (from LOO results)"],
    ["Szafrański/Joshi 2025 clinical",
     "N=127 cross-sectional, 5 genera subset; "
     "metadata: diagnosis (PIH/PIM/PI), GDI value",
     "Available (Datasets/20260416_mSystems_16S_5genera_all_profiles.txt)"],
]):
    for j, v in enumerate(row):
        tbl3.rows[i].cells[j].text = v
        if i == 0:
            tbl3.rows[i].cells[j].paragraphs[0].runs[0].bold = True
doc.add_paragraph()
P("All tables will be formatted for R/vegan (Bray-Curtis + PCoA or PERMANOVA) or "
  "Python/scikit-bio. I will send these as a follow-up. "
  "The PCA/LDA ordination already done in the analysis (figures fig_pca_pred_obs.png, "
  "fig4_b_pca.png) shows: (1) predicted compositions span the same data manifold as "
  "observed, (2) CT1/CT2 discrimination preserved under LOO (90%, only Patient A "
  "misclassified), (3) patient-specific b̂ PCA separates CT1/CT2 without supervision.")

# ── Action items ─────────────────────────────────────────
H("Summary of action items")
tbl4 = doc.add_table(rows=10, cols=3)
tbl4.style = "Table Grid"
for i, row in enumerate([
    ["#", "Owner", "Item"],
    ["1", "Keisuke", "Supplement PDF with model description sections (as above)"],
    ["2", "Keisuke", "Prepare intermediate data sheet: net_flow matrix F with "
                      "source citations for all 35 constrained pairs"],
    ["3", "Keisuke", "Document symmetric A assumption and its implications "
                      "for directionality (in methods section)"],
    ["4", "Keisuke", "Prepare LOO prediction tables for Dieckow (observed + predicted, "
                      "long format, PERMANOVA-ready)"],
    ["5", "Keisuke", "Send Szafrański 5-genera subset reformatted as feature+sample table"],
    ["6", "Keisuke", "Design A(pH,t) extension once WhatsApp sketches received"],
    ["7", "Szymon", "Send WhatsApp sketches on sucrose–lactate–pH scenarios"],
    ["8", "Szymon", "Share SimCom (Radek/Room Drum) results for comparison"],
    ["9", "Keisuke", "Review SimCom results and compare with Bacilli↔Negativicutes "
                      "muted interaction finding"],
]):
    for j, v in enumerate(row):
        tbl4.rows[i].cells[j].text = v
        if i == 0:
            tbl4.rows[i].cells[j].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
P("The finding that Bacilli↔Negativicutes (the Streptococcus–lactate–Veillonella axis "
  "you highlighted) is sign-consistent but ecologically muted in the 3-week data "
  "connects directly to your scenario: the sustained lactate concentration may simply "
  "not be reached at the abutment in the Dieckow recovery setting. "
  "SimCom data would be very valuable here — if it shows active lactate flux between "
  "these genera on shorter timescales, that would support the A(t) extension.")
doc.add_paragraph()
P("Best regards,")
P("Keisuke")

out = "/home/nishioka/IKM_Hiwi/docs/reply_szafranski_20260512.docx"
doc.save(out)
print(f"Saved: {out}")
